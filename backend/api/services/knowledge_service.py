"""
Knowledge Base Service
Handles document processing, embeddings, and vector search using Weaviate
"""

import os
import json
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib
import asyncio
from pathlib import Path

import weaviate
from weaviate.embedded import EmbeddedOptions
import openai
from openai import AsyncOpenAI
import PyPDF2
import docx
import tiktoken
from fastapi import UploadFile, HTTPException

from api.models.company import Company, KnowledgeBase, Document
from api.core.database import get_session_context
from sqlmodel import select
from api.core.config import get_settings

logger = logging.getLogger(__name__)

class KnowledgeService:
    """
    Service for managing company knowledge bases
    Handles document upload, processing, and retrieval
    """
    
    def __init__(self):
        # Initialize Weaviate client
        settings = get_settings()
        self.weaviate_client = None
        try:
            # v3 Client: positional URL
            self.weaviate_client = weaviate.Client(
                settings.WEAVIATE_URL,
                timeout_config=(5, 15)
            )
        except Exception as e:
            logger.warning(f"Weaviate client not initialized, falling back without vector store: {e}")
        
        # Initialize OpenAI client for embeddings (optional)
        self.openai_client = None
        try:
            if settings.OPENAI_API_KEY:
                self.openai_client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
            else:
                logger.info("OPENAI_API_KEY not set; embeddings and semantic search disabled")
        except Exception as e:
            logger.warning(f"OpenAI client not initialized; embeddings disabled: {e}")
        
        # Tokenizer for chunk size calculation
        self.encoding = tiktoken.encoding_for_model("gpt-4")
        
    async def create_knowledge_base(self, company: Company) -> KnowledgeBase:
        """
        Create a new knowledge base for a company
        """
        # Generate unique namespace
        namespace = f"Company_{company.id.replace('-', '_')}"
        
        # Create Weaviate schema for this company
        if self.weaviate_client:
            schema = {
                "class": namespace,
                "description": f"Knowledge base for {company.name}",
                "vectorizer": "text2vec-openai",
                "properties": [
                    {"name": "content", "dataType": ["text"], "description": "The content chunk"},
                    {"name": "document_id", "dataType": ["string"], "description": "Source document ID"},
                    {"name": "document_name", "dataType": ["string"], "description": "Source document name"},
                    {"name": "chunk_index", "dataType": ["int"], "description": "Position in document"},
                    {"name": "metadata", "dataType": ["text"], "description": "Additional metadata as JSON"},
                ],
            }
            try:
                self.weaviate_client.schema.create_class(schema)
            except Exception as e:
                if "already exists" not in str(e):
                    logger.error(f"Error creating Weaviate schema: {e}")
                    raise
        
        # Create database entry
        async with get_session_context() as session:
            knowledge_base = KnowledgeBase(
                company_id=company.id,
                weaviate_namespace=namespace
            )
            session.add(knowledge_base)
            await session.commit()
            await session.refresh(knowledge_base)
            
        return knowledge_base
    
    async def upload_document(
        self,
        knowledge_base: KnowledgeBase,
        file: UploadFile,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Upload and process a document
        """
        # Validate file type
        allowed_types = ['.pdf', '.docx', '.txt', '.md']
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in allowed_types:
            raise HTTPException(400, f"File type {file_ext} not supported")
        
        # Read file content
        content = await file.read()
        file_size = len(content)
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text = await self._extract_pdf_text(content)
        elif file_ext == '.docx':
            text = await self._extract_docx_text(content)
        else:
            text = content.decode('utf-8')
        
        # Create document record
        async with get_session_context() as session:
            document = Document(
                knowledge_base_id=knowledge_base.id,
                filename=file.filename,
                file_type=file_ext[1:],  # Remove dot
                file_size=file_size,
                original_text=text[:10000],  # Store first 10k chars for reference
                document_metadata=metadata or {},
                processing_status="processing"
            )
            session.add(document)
            await session.commit()
            await session.refresh(document)
        
        # Process document asynchronously
        asyncio.create_task(self._process_document(knowledge_base, document, text))
        
        return document
    
    async def _process_document(
        self,
        knowledge_base: KnowledgeBase,
        document: Document,
        text: str
    ):
        """
        Process document: chunk, embed, and store in Weaviate
        """
        try:
            # Chunk the text
            chunks = self._chunk_text(text, knowledge_base.chunk_size, knowledge_base.chunk_overlap)
            
            did_vectorize = False
            if self.weaviate_client and self.openai_client:
                try:
                    # Prepare batch for Weaviate
                    batch = []
                    for i, chunk in enumerate(chunks):
                        try:
                            embedding = await self._generate_embedding(chunk)
                        except Exception as embed_err:
                            logger.warning(f"Embedding failed, fallback to no-vector mode: {embed_err}")
                            batch = []
                            break
                        obj = {
                            "content": chunk,
                            "document_id": document.id,
                            "document_name": document.filename,
                            "chunk_index": i,
                            "metadata": json.dumps(document.document_metadata),
                        }
                        batch.append({
                            "class": knowledge_base.weaviate_namespace,
                            "properties": obj,
                            "vector": embedding,
                        })
                    if batch:
                        with self.weaviate_client.batch as batch_client:
                            for obj in batch:
                                batch_client.add_data_object(
                                    obj["properties"],
                                    obj["class"],
                                    vector=obj["vector"],
                                )
                        did_vectorize = True
                except Exception as ve:
                    logger.warning(f"Vector store operation failed, continuing without vectors: {ve}")
            
            # Update document status
            async with get_session_context() as session:
                document.chunk_count = len(chunks) if did_vectorize else 0
                document.processed_at = datetime.utcnow()
                document.processing_status = "completed"
                session.add(document)
                
                # Update knowledge base stats
                kb = await session.get(KnowledgeBase, knowledge_base.id)
                kb.document_count += 1
                kb.total_chunks += document.chunk_count
                kb.last_updated = datetime.utcnow()
                session.add(kb)
                
                await session.commit()
                
            logger.info(f"Successfully processed document {document.filename} with {document.chunk_count} chunks")
            
        except Exception as e:
            logger.error(f"Error processing document {document.id}: {e}")
            async with get_session_context() as session:
                document.processing_status = "failed"
                session.add(document)
                await session.commit()
    
    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split text into chunks based on token count
        """
        tokens = self.encoding.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), chunk_size - overlap):
            chunk_tokens = tokens[i:i + chunk_size]
            chunk_text = self.encoding.decode(chunk_tokens)
            chunks.append(chunk_text)
            
            if i + chunk_size >= len(tokens):
                break
        
        return chunks
    
    async def _generate_embedding(self, text: str) -> List[float]:
        """
        Generate embedding using OpenAI
        """
        if not self.openai_client:
            raise RuntimeError("Embeddings requested but OpenAI client is not configured")
        response = await self.openai_client.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
        return response.data[0].embedding
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """
        Extract text from PDF
        """
        import io
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """
        Extract text from DOCX
        """
        import io
        doc = docx.Document(io.BytesIO(content))
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text
    
    async def search_knowledge(
        self,
        knowledge_base: KnowledgeBase,
        query: str,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """
        Search knowledge base using semantic similarity
        Fallback: plain text search over stored document text when vector store is unavailable
        """
        # Vector search path
        if self.weaviate_client and self.openai_client:
            query_embedding = await self._generate_embedding(query)
            result = (
                self.weaviate_client.query
                .get(knowledge_base.weaviate_namespace, ["content", "document_name", "chunk_index", "metadata"])
                .with_near_vector({"vector": query_embedding})
                .with_limit(limit)
                .with_additional(["distance", "id"])
                .do()
            )
            if "data" not in result or "Get" not in result["data"]:
                return []
            chunks = result["data"]["Get"].get(knowledge_base.weaviate_namespace, [])
            formatted_results = []
            for chunk in chunks:
                formatted_results.append({
                    "content": chunk["content"],
                    "document": chunk["document_name"],
                    "chunk_index": chunk["chunk_index"],
                    "relevance_score": 1 - chunk["_additional"]["distance"],
                    "metadata": json.loads(chunk.get("metadata", "{}"))
                })
            return formatted_results

        # Fallback: plain text search over Document.original_text
        results: List[Dict[str, Any]] = []
        q_lower = query.strip().lower()
        if not q_lower:
            return []
        async with get_session_context() as session:
            from api.models.company import Document as KbDocument
            from sqlalchemy import select
            rows = await session.execute(
                select(KbDocument)
                .where(KbDocument.knowledge_base_id == knowledge_base.id)
            )
            docs = rows.scalars().all()
            for doc in docs:
                text = (doc.original_text or "")
                idx = text.lower().find(q_lower)
                if idx == -1:
                    continue
                # Create a simple snippet around the first match
                start = max(0, idx - 80)
                end = min(len(text), idx + len(query) + 80)
                snippet = text[start:end]
                # Score: inverse of index position and frequency
                freq = text.lower().count(q_lower)
                score = 1.0 / (1 + idx) + 0.1 * freq
                results.append({
                    "content": snippet,
                    "document": doc.filename,
                    "chunk_index": 0,
                    "relevance_score": round(min(score, 1.0), 4),
                    "metadata": {"document_id": doc.id, "match_index": idx, "occurrences": freq},
                })
        # Sort and limit
        results.sort(key=lambda r: r["relevance_score"], reverse=True)
        return results[:limit]
    
    async def delete_document(self, document: Document):
        """
        Delete a document and its chunks from Weaviate
        """
        # Get knowledge base
        async with get_session_context() as session:
            kb = await session.get(KnowledgeBase, document.knowledge_base_id)
        
        # Delete from Weaviate (if available)
        if self.weaviate_client:
            self.weaviate_client.batch.delete_objects(
                class_name=kb.weaviate_namespace,
                where={
                    "path": ["document_id"],
                    "operator": "Equal",
                    "valueString": document.id,
                },
            )
        
        # Update knowledge base stats
        async with get_session_context() as session:
            kb.document_count -= 1
            kb.total_chunks -= document.chunk_count
            kb.last_updated = datetime.utcnow()
            session.add(kb)
            
            # Delete document record
            await session.delete(document)
            await session.commit()
    
    async def get_knowledge_context(
        self,
        knowledge_base: KnowledgeBase,
        query: str,
        max_tokens: int = 2000
    ) -> str:
        """
        Get relevant context for a query (for RAG)
        """
        # Search for relevant chunks
        results = await self.search_knowledge(knowledge_base, query, limit=10)
        
        if not results:
            return ""
        
        # Combine chunks until we reach max tokens
        context = []
        token_count = 0
        
        for result in results:
            chunk_tokens = len(self.encoding.encode(result["content"]))
            if token_count + chunk_tokens > max_tokens:
                break
            
            context.append(f"[Source: {result['document']}]\n{result['content']}")
            token_count += chunk_tokens
        
        return "\n\n".join(context)